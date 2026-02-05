import os
import tempfile
from typing import List

import gradio as gr
import pandas as pd
import numpy as np

from dotenv import load_dotenv
from openai import AzureOpenAI
from azure.cosmos import CosmosClient

from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from rouge_score import rouge_scorer
from sacrebleu import sentence_bleu
from sklearn.metrics.pairwise import cosine_similarity

# --------------------------------------------------
# Load environment variables
# --------------------------------------------------
load_dotenv()

AZURE_OAI_ENDPOINT = os.getenv("AZURE_OAI_ENDPOINT")
AZURE_OAI_KEY = os.getenv("AZURE_OAI_KEY")
AZURE_OAI_DEPLOYMENT = os.getenv("AZURE_OAI_DEPLOYMENT")
AZURE_OAI_EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OAI_EMBEDDING_DEPLOYMENT")

COSMOS_ENDPOINT = os.getenv("COSMOS_ENDPOINT")
COSMOS_KEY = os.getenv("COSMOS_KEY")
COSMOS_DATABASE = os.getenv("COSMOS_DATABASE")
COSMOS_CONTAINER = os.getenv("COSMOS_CONTAINER")

# --------------------------------------------------
# Clients
# --------------------------------------------------
openai_client = AzureOpenAI(
    azure_endpoint=AZURE_OAI_ENDPOINT,
    api_key=AZURE_OAI_KEY,
    api_version="2024-02-15-preview"
)

cosmos_client = CosmosClient(COSMOS_ENDPOINT, COSMOS_KEY)
db = cosmos_client.get_database_client(COSMOS_DATABASE)
container = db.get_container_client(COSMOS_CONTAINER)

# --------------------------------------------------
# Utility functions
# --------------------------------------------------
def embed_text(text: str):
    if not text or not isinstance(text, str) or not text.strip():
        raise ValueError("Invalid text passed for embedding")

    return openai_client.embeddings.create(
        model=AZURE_OAI_EMBEDDING_DEPLOYMENT,
        input=text.strip()
    ).data[0].embedding


def retrieve_chunks(question: str, top_k: int = 5) -> List[str]:
    query_embedding = embed_text(question)

    query = """
    SELECT TOP @k c.content
    FROM c
    ORDER BY VectorDistance(c.embedding, @embedding)
    """

    results = container.query_items(
        query=query,
        parameters=[
            {"name": "@k", "value": top_k},
            {"name": "@embedding", "value": query_embedding}
        ],
        enable_cross_partition_query=True
    )

    return [r["content"] for r in results]


def compute_confidence_score(answer: str, true_answer: str):
    ans_emb = np.array(embed_text(answer)).reshape(1, -1)
    ref_emb = np.array(embed_text(true_answer)).reshape(1, -1)
    semantic_sim = cosine_similarity(ans_emb, ref_emb)[0][0]

    rouge = rouge_scorer.RougeScorer(["rougeL"], use_stemmer=True)
    rouge_l = rouge.score(true_answer, answer)["rougeL"].fmeasure

    score = round((0.65 * semantic_sim + 0.35 * rouge_l) * 100, 2)

    if score >= 85:
        return score, "High", "green"
    elif score >= 65:
        return score, "Medium", "orange"
    else:
        return score, "Low", "red"


def run_rag(question: str, user_instruction: str):
    chunks = retrieve_chunks(question)

    if not chunks:
        return {"answer": "No relevant knowledge found.", "score": 0, "label": "Low", "color": "red"}

    true_answer = chunks[0]

    # ---- SYSTEM PROMPT (AUTHORITATIVE, NEVER OVERRIDDEN) ----
    system_prompt = (
        "You are an information security assistant. "
        "Answer strictly using the provided context. "
        "Do not add external knowledge or assumptions."
    )

    # ---- USER PROMPT (OPTIONAL INSTRUCTIONS) ----
    user_prompt = f"""
Context:
{true_answer}

Question:
{question}

User Instructions (if any):
{user_instruction if user_instruction else "None"}
"""

    response = openai_client.chat.completions.create(
        model=AZURE_OAI_DEPLOYMENT,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.9,
        max_tokens=600
    )

    answer = response.choices[0].message.content.strip()
    score, label, color = compute_confidence_score(answer, true_answer)

    return {"answer": answer, "score": score, "label": label, "color": color}

# --------------------------------------------------
# File parsing
# --------------------------------------------------
def extract_questions(file_path: str) -> List[str]:
    if file_path.endswith(".csv"):
        df = pd.read_csv(file_path)
    elif file_path.endswith(".xlsx") or file_path.endswith(".xls"):
        df = pd.read_excel(file_path)
    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return [l.strip() for l in f if l.strip()]
    elif file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        text = "".join(p.extract_text() or "" for p in reader.pages)
        return [l.strip() for l in text.splitlines() if l.strip()]
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    else:
        raise ValueError("Unsupported file type")

    questions = []
    for col in df.columns:
        questions.extend(df[col].dropna().astype(str).tolist())

    return [q.strip() for q in questions if q.strip()]

# --------------------------------------------------
# Gradio processing
# --------------------------------------------------
def process_file(file, user_instruction):
    questions = extract_questions(file.name)

    if not questions:
        raise gr.Error("No questions found in the uploaded file.")

    rows = []

    for q in questions:
        result = run_rag(q, user_instruction)

        rows.append([
            q,
            result["answer"],
            result["score"],
            result["label"],
            ""  # User Feedback
        ])

    return pd.DataFrame(
        rows,
        columns=["Question", "Answer", "Confidence Score", "Label", "User Feedback"]
    )

# --------------------------------------------------
# Export utilities
# --------------------------------------------------
def export_csv(df):
    path = os.path.join(tempfile.gettempdir(), "Generated_Responses.csv")
    df.to_csv(path, index=False)
    return path


def export_txt(df):
    path = os.path.join(tempfile.gettempdir(), "Generated_Responses.txt")
    with open(path, "w", encoding="utf-8") as f:
        for _, row in df.iterrows():
            for col in df.columns:
                f.write(f"{col}: {row[col]}\n")
            f.write("\n" + "-" * 60 + "\n\n")
    return path


def export_docx(df):
    path = os.path.join(tempfile.gettempdir(), "Generated_Responses.docx")
    doc = Document()
    doc.add_heading("Generated Responses", level=1)

    for _, row in df.iterrows():
        for col in df.columns:
            doc.add_paragraph(f"{col}: {row[col]}")
        doc.add_page_break()

    doc.save(path)
    return path


def export_pdf(df):
    path = os.path.join(tempfile.gettempdir(), "Generated_Responses.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    y = A4[1] - 40

    for _, row in df.iterrows():
        text = c.beginText(40, y)
        text.setFont("Helvetica", 10)

        for col in df.columns:
            text.textLine(f"{col}: {row[col]}")
            y -= 14
            if y < 60:
                c.drawText(text)
                c.showPage()
                text = c.beginText(40, A4[1] - 40)
                text.setFont("Helvetica", 10)
                y = A4[1] - 40

        c.drawText(text)

    c.save()
    return path


def export_file(df, export_format):
    if export_format == "CSV":
        return export_csv(df)
    if export_format == "PDF":
        return export_pdf(df)
    if export_format == "Word":
        return export_docx(df)
    if export_format == "TXT":
        return export_txt(df)
    raise ValueError("Unsupported export format")

# --------------------------------------------------
# Custom CSS (adaptive light/dark)
# --------------------------------------------------
custom_css = """
.gradio-container { background: var(--background-fill-primary) !important; }
@media (prefers-color-scheme: light) {
  body { color: #000000; }
  tbody tr:hover { background: #e5e7eb !important; }
}
@media (prefers-color-scheme: dark) {
  body { color: #ffffff; }
  tbody tr:hover { background: #374151 !important; }
}
"""

# --------------------------------------------------
# Gradio UI
# --------------------------------------------------
with gr.Blocks(css=custom_css) as demo:

    gr.Markdown("<h1 style='text-align:center;'>Response Generation AI Agent</h1>")

    with gr.Row():
        file_input = gr.File(label="Upload Question File")
        user_prompt = gr.Textbox(
            label="Optional Answer Instructions",
            placeholder="e.g. Answer in not more than two sentences, max 50 words, bullet points, etc."
        )

    run_btn = gr.Button("Generate Responses")

    output_dataframe = gr.Dataframe(
        headers=["Question", "Answer", "Confidence Score", "Label", "User Feedback"],
        interactive=True,
        wrap=True
    )

    with gr.Row():
        export_format = gr.Dropdown(["CSV", "PDF", "Word", "TXT"], value="CSV")
        download_btn = gr.Button("Download Output")
        download_file = gr.File()

    run_btn.click(
        process_file,
        inputs=[file_input, user_prompt],
        outputs=output_dataframe
    )

    download_btn.click(
        export_file,
        inputs=[output_dataframe, export_format],
        outputs=download_file
    )

# --------------------------------------------------
# Launch
# --------------------------------------------------
if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=int(os.environ.get("PORT", 7860)))
